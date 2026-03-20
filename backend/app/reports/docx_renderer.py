from io import BytesIO

from docx import Document


def render_client_docx(context: dict) -> BytesIO:
    doc = Document()

    doc.add_heading("Отчет для клиента", level=1)
    doc.add_paragraph(f"Тест: {context['test_title']}")
    doc.add_paragraph(f"Участник: {context.get('participant_name') or '—'}")
    doc.add_paragraph(f"Email: {context.get('participant_email') or '—'}")
    doc.add_paragraph(
        f"Прогресс: {context['answered_count']} / {context['total_questions']} "
        f"({context['progress_percent']}%)"
    )

    doc.add_heading("Метрики", level=2)
    for scale in context["metrics"].get("scales", []):
        doc.add_paragraph(
            f"{scale['title']}: {scale['normalized_score']} "
            f"({scale['interpretation']})"
        )

    doc.add_heading("Краткий итог", level=2)
    summary = context["metrics"].get("summary", {})
    doc.add_paragraph(f"Количество шкал: {summary.get('total_scales', 0)}")
    doc.add_paragraph(f"Доминирующая шкала: {summary.get('dominant_scale') or '—'}")
    doc.add_paragraph(f"Значение: {summary.get('dominant_score', 0)}")

    stream = BytesIO()
    doc.save(stream)
    stream.seek(0)
    return stream


def render_pro_docx(context: dict) -> BytesIO:
    doc = Document()

    doc.add_heading("Отчет для профориентолога", level=1)
    doc.add_paragraph(f"Тест: {context['test_title']}")
    doc.add_paragraph(f"Участник: {context.get('participant_name') or '—'}")
    doc.add_paragraph(f"Email: {context.get('participant_email') or '—'}")
    doc.add_paragraph(f"Статус: {context['session_status']}")
    doc.add_paragraph(
        f"Прогресс: {context['answered_count']} / {context['total_questions']} "
        f"({context['progress_percent']}%)"
    )

    doc.add_heading("Метрики", level=2)
    for scale in context["metrics"].get("scales", []):
        doc.add_paragraph(
            f"{scale['title']} | raw={scale['raw_score']} | "
            f"norm={scale['normalized_score']} | "
            f"{scale['interpretation']}"
        )

    doc.add_heading("Ответы", level=2)
    for answer in context.get("answers", []):
        doc.add_paragraph(
            f"Вопрос {answer['question_id']}: "
            f"text={answer.get('value_text')} json={answer.get('value_json')}"
        )

    stream = BytesIO()
    doc.save(stream)
    stream.seek(0)
    return stream